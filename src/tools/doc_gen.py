import os
import json
import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
import mammoth
from langchain.tools import BaseTool
from langchain.pydantic_v1 import BaseModel, Field
from typing import Optional, List, Dict, Any
from crewai import Tool


class ExtractDocTemplateInput(BaseModel):
    document_path: str = Field(..., description="Path to the DOCX document to extract template from")


class WebScrapeInput(BaseModel):
    url: str = Field(..., description="URL of the website to scrape")
    selectors: Optional[List[str]] = Field(None, description="CSS selectors to target specific content")


class DocumentTools:
    @staticmethod
    def extract_doc_template_structure():
        """Creates a tool that extracts template structure from a DOCX document"""
        
        def _extract_doc_template(document_path: str) -> Dict[str, Any]:
            """Extract template structure from a DOCX document"""
            try:
                # Check if file exists
                if not os.path.exists(document_path):
                    return {"error": f"File not found: {document_path}"}
                
                # Read the document
                doc = Document(document_path)
                
                # Extract structure
                structure = {
                    "title": doc.core_properties.title or "Untitled Document",
                    "sections": [],
                    "styles": []
                }
                
                # Extract sections and content
                current_section = None
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                        
                    # Check if paragraph is a heading
                    if para.style.name.startswith('Heading'):
                        if current_section:
                            structure["sections"].append(current_section)
                            
                        level = int(para.style.name.replace('Heading ', '')) if para.style.name != 'Heading' else 1
                        current_section = {
                            "title": para.text,
                            "level": level,
                            "content": [],
                            "placeholders": []
                        }
                    elif current_section:
                        current_section["content"].append(para.text)
                        
                        # Identify potential placeholders (text in brackets)
                        import re
                        placeholders = re.findall(r'\{([^}]+)\}|\[([^]]+)\]', para.text)
                        if placeholders:
                            for ph in placeholders:
                                placeholder = next((p for p in ph if p), None)
                                if placeholder and placeholder not in current_section["placeholders"]:
                                    current_section["placeholders"].append(placeholder)
                    
                    # Track styles used
                    if para.style.name not in [s["name"] for s in structure["styles"]]:
                        structure["styles"].append({
                            "name": para.style.name,
                            "font": para.style.font.name if hasattr(para.style, 'font') and para.style.font else None,
                            "size": para.style.font.size if hasattr(para.style, 'font') and para.style.font else None
                        })
                
                # Add the last section if exists
                if current_section:
                    structure["sections"].append(current_section)
                
                # Extract tables
                structure["tables"] = []
                for i, table in enumerate(doc.tables):
                    table_data = {
                        "index": i,
                        "rows": len(table.rows),
                        "columns": len(table.columns) if table.rows else 0,
                        "headers": []
                    }
                    
                    # Extract headers (assuming first row is header)
                    if table.rows:
                        for cell in table.rows[0].cells:
                            table_data["headers"].append(cell.text.strip())
                    
                    structure["tables"].append(table_data)
                
                return structure
                
            except Exception as e:
                return {"error": str(e)}
        
        return Tool(
            name="extract_doc_template_structure",
            func=_extract_doc_template,
            description="Extracts the template structure from a DOCX document"
        )
    
    @staticmethod
    def convert_doc_to_text():
        """Creates a tool that converts a DOCX document to plain text"""
        
        def _convert_doc_to_text(document_path: str) -> str:
            """Convert a DOCX document to plain text"""
            try:
                # Check if file exists
                if not os.path.exists(document_path):
                    return f"Error: File not found: {document_path}"
                
                # Use mammoth to convert to text
                with open(document_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text = result.value
                
                return text
                
            except Exception as e:
                return f"Error converting document: {str(e)}"
        
        return Tool(
            name="convert_doc_to_text",
            func=_convert_doc_to_text,
            description="Converts a DOCX document to plain text"
        )
    
    @staticmethod
    def web_scrape_content():
        """Creates a tool that scrapes content from a website"""
        
        def _web_scrape(url: str, selectors: Optional[List[str]] = None) -> Dict[str, Any]:
            """Scrape content from a website"""
            try:
                # Send request
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = requests.get(url, headers=headers)
                response.raise_for_status()
                
                # Parse HTML
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Initialize result
                result = {
                    "url": url,
                    "title": soup.title.text if soup.title else "No title",
                    "content": {}
                }
                
                # Extract content based on selectors if provided
                if selectors:
                    for selector in selectors:
                        elements = soup.select(selector)
                        result["content"][selector] = [elem.get_text(strip=True) for elem in elements]
                else:
                    # Default extraction
                    # Get main content (heuristic)
                    main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
                    
                    if main_content:
                        result["content"]["main"] = main_content.get_text(strip=True)
                    
                    # Get paragraphs
                    paragraphs = soup.find_all('p')
                    result["content"]["paragraphs"] = [p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)]
                    
                    # Get headers
                    headers = soup.find_all(['h1', 'h2', 'h3'])
                    result["content"]["headers"] = [h.get_text(strip=True) for h in headers if h.get_text(strip=True)]
                
                return result
                
            except Exception as e:
                return {"error": str(e)}
        
        return Tool(
            name="web_scrape_content",
            func=_web_scrape,
            description="Scrapes content from a website using optional CSS selectors"
        )
    
    @staticmethod
    def generate_document_template():
        """Creates a tool that generates a document template based on requirements"""
        
        def _generate_template(document_type: str, sections: List[str], placeholders: Dict[str, str]) -> Dict[str, Any]:
            """Generate a document template based on requirements"""
            try:
                template = {
                    "document_type": document_type,
                    "template_structure": {
                        "title": f"{document_type.capitalize()} Template",
                        "sections": []
                    },
                    "placeholders": placeholders
                }
                
                # Create sections
                for i, section in enumerate(sections):
                    section_obj = {
                        "id": f"section_{i+1}",
                        "title": section,
                        "content": f"{{content_for_{section.lower().replace(' ', '_')}}}",
                        "placeholders": []
                    }
                    
                    # Add relevant placeholders to each section
                    for key, desc in placeholders.items():
                        if key.lower() in section.lower() or section.lower() in key.lower():
                            section_obj["placeholders"].append({
                                "key": key,
                                "description": desc
                            })
                    
                    template["template_structure"]["sections"].append(section_obj)
                
                return template
                
            except Exception as e:
                return {"error": str(e)}
        
        return Tool(
            name="generate_document_template",
            func=_generate_template,
            description="Generates a document template based on document type, sections, and placeholders"
        )
    
    @staticmethod
    def evaluate_document_quality():
        """Creates a tool that evaluates the quality of a document"""
        
        def _evaluate_document(document_text: str, criteria: Optional[Dict[str, float]] = None) -> Dict[str, Any]:
            """Evaluate the quality of a document based on specified criteria"""
            try:
                # Default criteria with weights if not provided
                if not criteria:
                    criteria = {
                        "coherence": 0.25,
                        "relevance": 0.25,
                        "grammar": 0.2,
                        "structure": 0.15,
                        "formatting": 0.15
                    }
                
                # Basic metrics
                word_count = len(document_text.split())
                sentence_count = document_text.count('.') + document_text.count('!') + document_text.count('?')
                avg_words_per_sentence = word_count / max(1, sentence_count)
                
                # Simple quality metrics (in practice, these would be much more sophisticated)
                metrics = {
                    "word_count": word_count,
                    "sentence_count": sentence_count,
                    "avg_words_per_sentence": avg_words_per_sentence,
                    "scores": {
                        # These would be calculated with much more sophisticated algorithms in practice
                        "coherence": min(0.95, 0.8 + (50 < avg_words_per_sentence < 25) * 0.15),
                        "relevance": 0.85,  # Placeholder - would require content analysis
                        "grammar": 0.9,     # Placeholder - would require grammar checking
                        "structure": 0.85 if document_text.count('\n\n') > 3 else 0.7,
                        "formatting": 0.8   # Placeholder
                    },
                    "feedback": []
                }
                
                # Calculate weighted score
                weighted_score = 0
                for criterion, weight in criteria.items():
                    if criterion in metrics["scores"]:
                        weighted_score += metrics["scores"][criterion] * weight
                
                metrics["overall_score"] = weighted_score
                
                # Generate feedback (simplified - would be more detailed in practice)
                if avg_words_per_sentence > 30:
                    metrics["feedback"].append("Consider shortening sentences for better readability.")
                if word_count < 300:
                    metrics["feedback"].append("Document may be too short for comprehensive coverage.")
                if document_text.count('\n\n') < 5:
                    metrics["feedback"].append("Consider adding more paragraph breaks for better structure.")
                
                return metrics
                
            except Exception as e:
                return {"error": str(e)}
        
        return Tool(
            name="evaluate_document_quality",
            func=_evaluate_document,
            description="Evaluates the quality of a document based on various criteria"
        )
    
    @staticmethod
    def create_all_tools():
        """Creates all document generation tools"""
        return [
            DocumentTools.extract_doc_template_structure(),
            DocumentTools.convert_doc_to_text(),
            DocumentTools.web_scrape_content(),
            DocumentTools.generate_document_template(),
            DocumentTools.evaluate_document_quality()
        ]


# Example usage
if __name__ == "__main__":
    tools = DocumentTools.create_all_tools()
    
    # Example: Extract template from a document
    # result = tools[0].func("example.docx")
    # print(json.dumps(result, indent=2))