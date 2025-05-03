"""
DOC Template Structure Extractor

This script extracts the template structure from a DOC file and converts it to JSON,
capturing text areas, images, headings, fonts, and positioning information for each page.
"""

import os
import json
import zipfile
import re
import io
import xml.etree.ElementTree as ET
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import sys
import datetime

# Namespaces used in Word OOXML files
namespaces = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'v': 'urn:schemas-microsoft-com:vml',
    'o': 'urn:schemas-microsoft-com:office:office',
}

def extract_doc_template_structure(doc_path):
    """
    Extract template structure from a DOC/DOCX file
    
    Args:
        doc_path (str): Path to the DOC/DOCX file
        
    Returns:
        dict: JSON representation of the template structure
    """
    try:
        # Load the document
        doc = docx.Document(doc_path)
        
        # If it's a .doc file, convert to .docx first
        if doc_path.lower().endswith('.doc'):
            temp_docx_path = doc_path + 'x'
            doc.save(temp_docx_path)
            doc_path = temp_docx_path
            doc = docx.Document(doc_path)
        
        # Extract document properties
        template_structure = {
            "metadata": {
                "filename": os.path.basename(doc_path),
                "extractedAt": datetime.datetime.now().isoformat(),
                "pageSize": extract_page_size(doc),
                "margins": extract_margins(doc),
                "totalPages": count_pages(doc_path)
            },
            "pages": extract_pages(doc_path, doc)
        }
        
        # Remove temporary file if created
        if doc_path.endswith('x') and doc_path != doc_path[:-1]:
            try:
                os.remove(doc_path)
            except:
                pass
        
        return template_structure
    
    except Exception as e:
        print(f"Error extracting template structure: {str(e)}")
        raise

def extract_page_size(doc):
    """Extract the page size from the document"""
    sections = doc.sections
    if not sections:
        return {"width": "8.5in", "height": "11in"}  # Default letter size
    
    section = sections[0]
    return {
        "width": f"{section.page_width.inches:.2f}in",
        "height": f"{section.page_height.inches:.2f}in"
    }

def extract_margins(doc):
    """Extract margin information from the document"""
    sections = doc.sections
    if not sections:
        return {"top": "1in", "right": "1in", "bottom": "1in", "left": "1in"}
    
    section = sections[0]
    return {
        "top": f"{section.top_margin.inches:.2f}in",
        "right": f"{section.right_margin.inches:.2f}in",
        "bottom": f"{section.bottom_margin.inches:.2f}in",
        "left": f"{section.left_margin.inches:.2f}in"
    }

def count_pages(docx_path):
    """Count the number of pages in the document"""
    try:
        # Extract the document.xml content from the DOCX file
        with zipfile.ZipFile(docx_path, 'r') as z:
            # Check if the document has the word/document.xml file
            if 'word/document.xml' in z.namelist():
                xml_content = z.read('word/document.xml')
                
                # Count page breaks
                root = ET.fromstring(xml_content)
                page_breaks = root.findall('.//w:lastRenderedPageBreak', namespaces)
                section_breaks = root.findall('.//w:sectPr', namespaces)
                
                # Each section break typically creates a new page
                # Add 1 because the document starts with at least one page
                return len(page_breaks) + len(section_breaks) + 1
    except Exception as e:
        print(f"Error counting pages: {str(e)}")
    
    # Fallback to estimating based on paragraphs (rough estimate)
    doc = docx.Document(docx_path)
    paragraph_count = len(doc.paragraphs)
    return max(1, paragraph_count // 30)  # Roughly 30 paragraphs per page

def extract_pages(docx_path, doc):
    """Extract content for each page"""
    pages = []
    
    # Open the docx as a zip file to access the XML directly
    with zipfile.ZipFile(docx_path) as docx_zip:
        # Extract document.xml
        xml_content = docx_zip.read('word/document.xml')
        root = ET.fromstring(xml_content)
        
        # Extract relationships to get image information
        rels_xml = docx_zip.read('word/_rels/document.xml.rels')
        rels_root = ET.fromstring(rels_xml)
        relationships = {rel.get('Id'): rel.get('Target') for rel in rels_root}
        
        # Parse sections to determine page breaks
        sections = get_section_breaks(root)
        
        # Process each paragraph and run for content and styles
        content_elements = []
        
        # Get body element that contains all content
        body = root.find('.//w:body', namespaces)
        
        # Process all elements in the body
        current_page = 1
        page_elements = []
        
        for elem in body:
            # Check if this element is a paragraph
            if elem.tag == f"{{{namespaces['w']}}}p":
                paragraph_data = extract_paragraph_info(elem, doc)
                
                # Check if this paragraph has a page break
                if has_page_break(elem):
                    # Save current page and start a new one
                    pages.append({
                        "pageNumber": current_page,
                        "elements": page_elements
                    })
                    current_page += 1
                    page_elements = []
                
                page_elements.append(paragraph_data)
                
                # Process runs within the paragraph for images
                for run in elem.findall('.//w:r', namespaces):
                    drawing = run.find('.//w:drawing', namespaces)
                    if drawing is not None:
                        image_data = extract_image_info(drawing, relationships, docx_zip)
                        if image_data:
                            page_elements.append(image_data)
            
            # Check for tables
            elif elem.tag == f"{{{namespaces['w']}}}tbl":
                table_data = extract_table_info(elem, doc)
                page_elements.append(table_data)
        
        # Add the last page
        if page_elements:
            pages.append({
                "pageNumber": current_page,
                "elements": page_elements
            })
    
    return pages

def get_section_breaks(root):
    """Get all section breaks in the document"""
    sections = []
    for sectPr in root.findall('.//w:sectPr', namespaces):
        section = {}
        
        # Get page size
        page_size = sectPr.find('.//w:pgSz', namespaces)
        if page_size is not None:
            section['width'] = page_size.get(f"{{{namespaces['w']}}}w")
            section['height'] = page_size.get(f"{{{namespaces['w']}}}h")
        
        # Get margins
        page_mar = sectPr.find('.//w:pgMar', namespaces)
        if page_mar is not None:
            section['margins'] = {
                'top': page_mar.get(f"{{{namespaces['w']}}}top"),
                'right': page_mar.get(f"{{{namespaces['w']}}}right"),
                'bottom': page_mar.get(f"{{{namespaces['w']}}}bottom"),
                'left': page_mar.get(f"{{{namespaces['w']}}}left")
            }
        
        sections.append(section)
    
    return sections

def has_page_break(paragraph_elem):
    """Check if a paragraph contains a page break"""
    # Check for explicit page breaks
    br_elements = paragraph_elem.findall('.//w:br', namespaces)
    for br in br_elements:
        break_type = br.get(f"{{{namespaces['w']}}}type")
        if break_type == 'page':
            return True
    
    # Check for section breaks (which create page breaks)
    if paragraph_elem.find('.//w:sectPr', namespaces) is not None:
        return True
    
    return False

def extract_paragraph_info(paragraph_elem, doc):
    """Extract information about a paragraph"""
    paragraph_style = {}
    
    # Get style name
    style_elem = paragraph_elem.find('.//w:pStyle', namespaces)
    if style_elem is not None:
        paragraph_style['styleName'] = style_elem.get(f"{{{namespaces['w']}}}val")
    
    # Get alignment
    alignment_elem = paragraph_elem.find('.//w:jc', namespaces)
    if alignment_elem is not None:
        alignment_val = alignment_elem.get(f"{{{namespaces['w']}}}val")
        if alignment_val:
            paragraph_style['alignment'] = alignment_val
    
    # Get paragraph indentation
    ind_elem = paragraph_elem.find('.//w:ind', namespaces)
    if ind_elem is not None:
        paragraph_style['indentation'] = {
            'left': ind_elem.get(f"{{{namespaces['w']}}}left", "0"),
            'right': ind_elem.get(f"{{{namespaces['w']}}}right", "0"),
            'firstLine': ind_elem.get(f"{{{namespaces['w']}}}firstLine", "0"),
            'hanging': ind_elem.get(f"{{{namespaces['w']}}}hanging", "0")
        }
    
    # Get spacing
    spacing_elem = paragraph_elem.find('.//w:spacing', namespaces)
    if spacing_elem is not None:
        paragraph_style['spacing'] = {
            'before': spacing_elem.get(f"{{{namespaces['w']}}}before", "0"),
            'after': spacing_elem.get(f"{{{namespaces['w']}}}after", "0"),
            'line': spacing_elem.get(f"{{{namespaces['w']}}}line", "240"),  # Default single spacing
            'lineRule': spacing_elem.get(f"{{{namespaces['w']}}}lineRule", "auto")
        }
    
    # Extract text content and runs (with their formatting)
    runs = []
    for run_elem in paragraph_elem.findall('.//w:r', namespaces):
        run_data = extract_run_info(run_elem)
        if run_data.get('text'):  # Only add runs with actual text
            runs.append(run_data)
    
    # Create paragraph data structure
    return {
        "type": "paragraph",
        "style": paragraph_style,
        "textRuns": runs,
        "text": ''.join([run.get('text', '') for run in runs])
    }

def extract_run_info(run_elem):
    """Extract formatting information from a text run"""
    run_data = {
        "text": "",
        "style": {}
    }
    
    # Get text content
    text_elements = run_elem.findall('.//w:t', namespaces)
    for t in text_elements:
        run_data["text"] += t.text if t.text else ""
    
    # Check for bold
    bold = run_elem.find('.//w:b', namespaces)
    if bold is not None:
        run_data["style"]["bold"] = True
    
    # Check for italic
    italic = run_elem.find('.//w:i', namespaces)
    if italic is not None:
        run_data["style"]["italic"] = True
    
    # Check for underline
    underline = run_elem.find('.//w:u', namespaces)
    if underline is not None:
        run_data["style"]["underline"] = True
    
    # Get font properties
    font = run_elem.find('.//w:rFonts', namespaces)
    if font is not None:
        run_data["style"]["font"] = font.get(f"{{{namespaces['w']}}}ascii") or font.get(f"{{{namespaces['w']}}}hAnsi")
    
    # Get size
    size = run_elem.find('.//w:sz', namespaces)
    if size is not None:
        # Size is in half-points, divide by 2 to get points
        run_data["style"]["size"] = f"{int(size.get(f'{{{namespaces['w']}}}val', '24')) / 2}pt"
    
    # Get color
    color = run_elem.find('.//w:color', namespaces)
    if color is not None:
        run_data["style"]["color"] = color.get(f"{{{namespaces['w']}}}val")
    
    # Get highlight
    highlight = run_elem.find('.//w:highlight', namespaces)
    if highlight is not None:
        run_data["style"]["highlight"] = highlight.get(f"{{{namespaces['w']}}}val")
    
    return run_data

def extract_image_info(drawing_elem, relationships, docx_zip):
    """Extract information about an image"""
    try:
        # Find the blip (image reference)
        blip = drawing_elem.find('.//a:blip', namespaces)
        if blip is None:
            return None
        
        # Get the relationship ID for the image
        rel_id = blip.get(f"{{{namespaces['r']}}}embed")
        if not rel_id or rel_id not in relationships:
            return None
        
        # Get positioning information
        extent = drawing_elem.find('.//wp:extent', namespaces)
        position = drawing_elem.find('.//wp:positionH', namespaces) or drawing_elem.find('.//wp:posOffset', namespaces)
            
        # Get image path within the docx
        image_path = f"word/{relationships[rel_id]}"
        
        # Extract image dimensions from the XML
        width = height = "0"
        if extent is not None:
            width = extent.get('cx')
            height = extent.get('cy')
            # Convert EMUs (English Metric Units) to inches (1 inch = 914400 EMUs)
            width_in = f"{float(width) / 914400:.2f}in"
            height_in = f"{float(height) / 914400:.2f}in"
        
        # Get positioning information
        positioning = {}
        if position is not None:
            # Try to get the positioning type and value
            rel_from = position.get(f"{{{namespaces['wp']}}}relativeFrom")
            if rel_from:
                positioning["relativeFrom"] = rel_from
            
            # Try to get offset
            pos_offset = position.find('.//wp:posOffset', namespaces)
            if pos_offset is not None and pos_offset.text:
                # Convert EMUs to inches
                positioning["offset"] = f"{float(pos_offset.text) / 914400:.2f}in"
        
        # Try to get content type
        content_type = "image/jpeg"  # Default
        if image_path.lower().endswith('.png'):
            content_type = "image/png"
        elif image_path.lower().endswith('.gif'):
            content_type = "image/gif"
        
        return {
            "type": "image",
            "width": width_in if 'width_in' in locals() else "unknown",
            "height": height_in if 'height_in' in locals() else "unknown",
            "positioning": positioning,
            "contentType": content_type,
            "path": image_path
        }
    
    except Exception as e:
        print(f"Error extracting image info: {str(e)}")
        return None

def extract_table_info(table_elem, doc):
    """Extract information about a table"""
    table_data = {
        "type": "table",
        "properties": {},
        "rows": []
    }
    
    # Get table properties
    tbl_pr = table_elem.find('.//w:tblPr', namespaces)
    if tbl_pr is not None:
        # Get table width
        tbl_w = tbl_pr.find('.//w:tblW', namespaces)
        if tbl_w is not None:
            width_type = tbl_w.get(f"{{{namespaces['w']}}}type")
            width_val = tbl_w.get(f"{{{namespaces['w']}}}w")
            
            if width_type == 'pct':
                # Convert percentage value (50 = 50%)
                table_data["properties"]["width"] = f"{int(width_val) / 50}%"
            elif width_type == 'dxa':
                # Convert twentieths of a point to inches (1440 = 1 inch)
                table_data["properties"]["width"] = f"{int(width_val) / 1440:.2f}in"
        
        # Get table alignment
        jc = tbl_pr.find('.//w:jc', namespaces)
        if jc is not None:
            table_data["properties"]["alignment"] = jc.get(f"{{{namespaces['w']}}}val")
    
    # Get row information
    rows = table_elem.findall('.//w:tr', namespaces)
    for row_index, row in enumerate(rows):
        row_data = {"cells": []}
        
        # Get cells in this row
        cells = row.findall('.//w:tc', namespaces)
        for cell_index, cell in enumerate(cells):
            # Get cell content (paragraphs)
            cell_paragraphs = []
            for p in cell.findall('.//w:p', namespaces):
                paragraph_data = extract_paragraph_info(p, doc)
                cell_paragraphs.append(paragraph_data)
            
            # Get cell properties
            properties = {}
            
            # Check for merged cells
            grid_span = cell.find('.//w:gridSpan', namespaces)
            if grid_span is not None:
                properties["gridSpan"] = grid_span.get(f"{{{namespaces['w']}}}val")
            
            # Check for vertical merge
            v_merge = cell.find('.//w:vMerge', namespaces)
            if v_merge is not None:
                val = v_merge.get(f"{{{namespaces['w']}}}val", "continue")
                properties["verticalMerge"] = val
            
            # Add cell data
            row_data["cells"].append({
                "content": cell_paragraphs,
                "properties": properties
            })
        
        table_data["rows"].append(row_data)
    
    return table_data

def save_template_structure(template_structure, output_path=None):
    """
    Save the template structure as a JSON file
    
    Args:
        template_structure (dict): Template structure data
        output_path (str, optional): Path to save the JSON file. If None, derives from input file.
        
    Returns:
        str: Path to the saved JSON file
    """
    if output_path is None:
        output_path = f"template_structure_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.json"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(template_structure, f, indent=2, ensure_ascii=False)
    
    return output_path

def main():
    """Main function to run the script from command line"""
    if len(sys.argv) < 2:
        print("Usage: python doc_template_extractor.py <path_to_doc_file> [output_json_path]")
        return
    
    doc_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        print(f"Extracting template structure from {doc_path}...")
        template_structure = extract_doc_template_structure(doc_path)
        
        json_path = save_template_structure(template_structure, output_path)
        print(f"Template structure saved to {json_path}")
        
    except Exception as e:
        print(f"Error: {str(e)}")
        raise

if __name__ == "__main__":
    main()