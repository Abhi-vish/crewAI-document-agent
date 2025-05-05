import streamlit as st
import os
import tempfile
from pathlib import Path
import base64
import time
from crewai import Agent, Task, Crew, Process
import docx
import PyPDF2
import mammoth
import magic

# Import local modules
from src.config.agents import create_research_agent, create_template_analyzer, create_content_generator, create_document_assembler
from src.config.tasks import create_tasks

class DocumentTransformer:
    def __init__(self):
        # Create agents
        self.research_agent = create_research_agent()
        self.template_analyzer = create_template_analyzer()
        self.content_generator = create_content_generator()
        self.document_assembler = create_document_assembler()

        # Create tasks
        self.tasks = create_tasks(
            self.research_agent,
            self.template_analyzer,
            self.content_generator,
            self.document_assembler
        )

        # Setup crew
        self.crew = Crew(
            agents=[
                self.research_agent,
                self.template_analyzer,
                self.content_generator,
                self.document_assembler
            ],
            tasks=self.tasks,
            process=Process.sequential,
            verbose=True
        )

    def transform_document(self, template, query):
        # Update task descriptions
        self.crew.tasks[0].description += f"QUERY:\n{query}\n\nFind current, accurate information relevant to this query that will help create a compelling document."
        self.crew.tasks[1].description += f"\n\nTEMPLATE:\n{template}\n\nQUERY:\n{query}"
        self.crew.tasks[2].description += f"\n\nQUERY:\n{query}"
        self.crew.tasks[3].description += f"\n\nORIGINAL TEMPLATE:\n{template}"

        # Run crew
        result = self.crew.kickoff()
        return result

def extract_text_from_file(uploaded_file):
    """
    Extract text from various file formats uploaded through Streamlit
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as temp_file:
        temp_file.write(uploaded_file.getvalue())
        temp_path = temp_file.name
    
    # Detect file type using file extension
    file_ext = uploaded_file.name.split('.')[-1].lower()
    
    try:
        # Extract text based on file type
        if file_ext == 'docx':
            # Process DOCX
            with open(temp_path, 'rb') as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text = result.value
        elif file_ext == 'pdf':
            # Process PDF
            text = ""
            with open(temp_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        else:  # Default to text file
            # Process TXT/MD
            with open(temp_path, "r", encoding="utf-8") as text_file:
                text = text_file.read()
    finally:
        # Clean up the temporary file
        os.unlink(temp_path)
    
    return text

def generate_docx(content):
    """Generate a DOCX file from the content and return as bytes"""
    doc = docx.Document()
    
    # Split content by lines and add to document
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip():  # Skip empty lines
            # Check if paragraph is a heading (starts with #)
            if para.startswith('# '):
                doc.add_heading(para[2:], 0)  # Title
            elif para.startswith('## '):
                doc.add_heading(para[3:], 1)
            elif para.startswith('### '):
                doc.add_heading(para[4:], 2)
            elif para.startswith('#### '):
                doc.add_heading(para[5:], 3)
            else:
                doc.add_paragraph(para)
    
    # Save to bytes
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    
    with open(tmp_path, 'rb') as f:
        docx_bytes = f.read()
    
    # Clean up
    os.unlink(tmp_path)
    
    return docx_bytes

def get_download_link(content, filename, file_format):
    """Generate a download link for the content"""
    if file_format == 'docx':
        file_bytes = generate_docx(content)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:  # Default to txt
        file_bytes = content.encode()
        mime_type = "text/plain"
    
    b64 = base64.b64encode(file_bytes).decode()
    href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}">Download {file_format.upper()} File</a>'
    return href

def streamlit_ui():
    st.set_page_config(
        page_title="Document Transformer",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("🔄 AI Document Transformer")
    st.markdown("""
    Transform your document templates with AI assistance. Upload a template document and provide a query, 
    and our AI agents will research, analyze, and transform it into a compelling new document.
    """)
    
    # Sidebar
    st.sidebar.header("About")
    st.sidebar.info("""
    This application uses multiple AI agents to transform document templates:
    1. **Research Agent** finds relevant information
    2. **Template Analyzer** analyzes your document structure
    3. **Content Generator** creates new content
    4. **Document Assembler** puts everything together
    """)
    
    st.sidebar.header("Settings")
    output_format = st.sidebar.selectbox("Output Format", ["txt", "docx"], index=1)
    
    # Main UI
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("Upload Template")
        uploaded_file = st.file_uploader("Choose a template file", type=["txt", "docx", "pdf", "md"])
        
    with col2:
        st.subheader("Transformation Query")
        query = st.text_area("Enter your query or instructions", height=150)
    
    # Initialize session state
    if 'template_content' not in st.session_state:
        st.session_state.template_content = None
    if 'result' not in st.session_state:
        st.session_state.result = None
    if 'is_processing' not in st.session_state:
        st.session_state.is_processing = False
    
    # Process uploaded file
    if uploaded_file is not None:
        try:
            st.session_state.template_content = extract_text_from_file(uploaded_file)
            with st.expander("Preview Template Content"):
                st.text_area("Template Content", st.session_state.template_content, height=300)
        except Exception as e:
            st.error(f"Error processing the file: {str(e)}")
    
    # Process button
    if st.button("Transform Document", disabled=st.session_state.is_processing or not query or st.session_state.template_content is None):
        if not st.session_state.template_content:
            st.warning("Please upload a template document first.")
        elif not query:
            st.warning("Please enter a transformation query.")
        else:
            st.session_state.is_processing = True
            
            # Create a progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                status_text.text("Initializing transformation process...")
                progress_bar.progress(10)
                time.sleep(0.5)
                
                # Initialize the transformer
                transformer = DocumentTransformer()
                
                status_text.text("Research agent is gathering information...")
                progress_bar.progress(25)
                time.sleep(0.5)
                
                status_text.text("Template analyzer is analyzing the document structure...")
                progress_bar.progress(40)
                time.sleep(0.5)
                
                status_text.text("Content generator is creating new content...")
                progress_bar.progress(60)
                time.sleep(0.5)
                
                status_text.text("Document assembler is putting everything together...")
                progress_bar.progress(80)
                time.sleep(0.5)
                
                # Transform the document
                st.session_state.result = transformer.transform_document(st.session_state.template_content, query)
                
                # Complete the progress
                progress_bar.progress(100)
                status_text.text("Transformation complete!")
                time.sleep(0.5)
                
                # Clear progress elements
                progress_bar.empty()
                status_text.empty()
                
                st.success("Document transformation completed successfully!")
            except Exception as e:
                st.error(f"Error during transformation: {str(e)}")
            finally:
                st.session_state.is_processing = False
    
    # Display result
    if st.session_state.result:
        st.subheader("Transformed Document")
        
        # Display tabs for preview and download
        tab1, tab2 = st.tabs(["Preview", "Download"])
        
        with tab1:
            st.markdown("### Transformed Content")
            st.write(st.session_state.result.raw)        
        with tab2:
            filename = f"transformed_document.{output_format}"
            download_link = get_download_link(
                str(st.session_state.result), 
                filename, 
                output_format
            )
            st.markdown(download_link, unsafe_allow_html=True)
            
            st.info(f"Click the link above to download your transformed document as a {output_format.upper()} file.")

if __name__ == "__main__":
    streamlit_ui()